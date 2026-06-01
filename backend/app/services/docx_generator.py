"""Генерация .docx документов с правильным оформлением."""
import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class DocxGenerator:
    """Генератор юридических документов в формате .docx."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Настройка стилей для юридических документов."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.5

        # Поля страницы (ГОСТ)
        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)

    def add_title(self, text: str):
        """Добавить заголовок (центрированный, жирный)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(0)

    def add_subtitle(self, text: str):
        """Добавить подзаголовок (центрированный)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(12)

    def add_body_text(self, text: str, indent: bool = True, bold: bool = False):
        """Добавить основной текст документа."""
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.bold = bold
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    def add_section_header(self, text: str):
        """Добавить заголовок раздела (жирный, по центру)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def add_address_block(self, lines: list, align_right: bool = True):
        """Добавить блок адресата (шапка документа)."""
        for line in lines:
            p = self.doc.add_paragraph()
            if align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    def add_list_item(self, text: str, number: str = None, indent: bool = True):
        """Добавить элемент списка."""
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        if number:
            run = p.add_run(f"{number}. {text}")
        else:
            run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(4)

    def add_signature_block(self, name: str, date_placeholder: bool = True):
        """Добавить блок подписи."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(name)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(0)

        if date_placeholder:
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(12)
            run2 = p2.add_run("«___» ____________ 20__ г.")
            run2.font.size = Pt(14)
            run2.font.name = 'Times New Roman'

    def add_page_break(self):
        """Добавить разрыв страницы."""
        self.doc.add_page_break()

    def _add_watermark(self):
        """Добавить водяной знак (футер) на каждую страницу."""
        for section in self.doc.sections:
            footer = section.footer
            # Если в футере уже есть параграф, используем его, иначе создаем
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Разделительная линия (опционально)
            run_line = p.add_run("________________________________________________\n")
            run_line.font.size = Pt(8)
            run_line.font.color.rgb = RGBColor(150, 150, 150)
            
            run = p.add_run("Документ составлен и проверен ИИ-юристом LAXLY AI LAW — laxly.ru")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(100, 100, 100)

    def generate_claim(self, data: Dict[str, Any]) -> bytes:
        """Сгенерировать исковое заявление."""
        # Шапка
        if data.get('court_name'):
            self.add_address_block([
                f"В {data['court_name']}",
                ""
            ])

        plaintiff = data.get('plaintiff', {})
        defendant = data.get('defendant', {})

        # Истец
        plaintiff_lines = ["Истец:"]
        if plaintiff.get('name'):
            plaintiff_lines.append(f"  {plaintiff['name']}")
        if plaintiff.get('inn'):
            plaintiff_lines.append(f"  ИНН: {plaintiff['inn']}")
        if plaintiff.get('address'):
            plaintiff_lines.append(f"  Адрес: {plaintiff['address']}")
        
        defendant_lines = ["Ответчик:"]
        if defendant.get('name'):
            defendant_lines.append(f"  {defendant['name']}")
        if defendant.get('inn'):
            defendant_lines.append(f"  ИНН: {defendant['inn']}")
        if defendant.get('address'):
            defendant_lines.append(f"  Адрес: {defendant['address']}")

        for line in plaintiff_lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

        for line in defendant_lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

        # Разделитель
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run("")
        run.font.size = Pt(6)

        # Заголовок
        self.add_title("Исковое заявление")

        # Обстоятельства дела
        if data.get('circumstances'):
            self.add_section_header("1. Обстоятельства дела")
            self.add_body_text(data['circumstances'])

        # Правовое обоснование
        if data.get('legal_basis'):
            self.add_section_header("2. Правовое обоснование")
            self.add_body_text(data['legal_basis'])

        # Требования
        if data.get('claims'):
            self.add_section_header("3. Требования")
            claims = data['claims'] if isinstance(data['claims'], list) else data['claims'].split('\n')
            for i, claim in enumerate(claims, 1):
                if claim.strip():
                    self.add_list_item(claim.strip(), number=str(i))

        # Приложения
        self.add_section_header("4. Приложения")
        self.add_list_item("Копия искового заявления по числу ответчиков")
        self.add_list_item("Документы, подтверждающие обстоятельства дела")
        self.add_list_item("Расчёт взыскиваемой суммы")
        self.add_list_item("Квитанция об уплате государственной пошлины")

        # Подпись
        if plaintiff.get('name'):
            self.add_signature_block(plaintiff['name'])

        # Добавляем водяной знак перед сохранением
        self._add_watermark()

        # Сохраняем в bytes
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_complaint(self, data: Dict[str, Any]) -> bytes:
        """Сгенерировать жалобу."""
        # Шапка
        if data.get('authority_name'):
            self.add_address_block([
                f"В {data['authority_name']}",
                ""
            ])

        applicant = data.get('applicant', {})
        applicant_lines = ["Заявитель:"]
        if applicant.get('name'):
            applicant_lines.append(f"  {applicant['name']}")
        if applicant.get('inn'):
            applicant_lines.append(f"  ИНН: {applicant['inn']}")
        if applicant.get('address'):
            applicant_lines.append(f"  Адрес: {applicant['address']}")

        for line in applicant_lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

        if data.get('interested_party'):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(f"Заинтересованное лицо: {data['interested_party']}")
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(6)

        self.add_title("Жалоба")

        if data.get('appealed_action'):
            self.add_section_header("1. Обжалуемое действие (решение)")
            self.add_body_text(data['appealed_action'])

        if data.get('grounds'):
            self.add_section_header("2. Основания жалобы")
            self.add_body_text(data['grounds'])

        if data.get('claims'):
            self.add_section_header("3. Требования")
            claims = data['claims'] if isinstance(data['claims'], list) else data['claims'].split('\n')
            for i, claim in enumerate(claims, 1):
                if claim.strip():
                    self.add_list_item(claim.strip(), number=str(i))

        if applicant.get('name'):
            self.add_signature_block(applicant['name'])

        # Добавляем водяной знак перед сохранением
        self._add_watermark()

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_demand(self, data: Dict[str, Any]) -> bytes:
        """Сгенерировать досудебную претензию."""
        demander = data.get('demander', {})
        demander_from = data.get('demander_from', {})

        # Кому
        if demander.get('name'):
            self.add_address_block([
                f"Кому: {demander['name']}",
                f"ИНН: {demander.get('inn', '')}" if demander.get('inn') else "",
                f"Адрес: {demander.get('address', '')}" if demander.get('address') else "",
            ])

        # Разделитель
        p = self.doc.add_paragraph()
        run = p.add_run("")
        run.font.size = Pt(6)

        # От кого
        if demander_from.get('name'):
            p = self.doc.add_paragraph()
            run = p.add_run(f"От: {demander_from['name']}")
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            if demander_from.get('inn'):
                p2 = self.doc.add_paragraph()
                run2 = p2.add_run(f"ИНН: {demander_from['inn']}")
                run2.font.size = Pt(14)
                run2.font.name = 'Times New Roman'
            if demander_from.get('address'):
                p3 = self.doc.add_paragraph()
                run3 = p3.add_run(f"Адрес: {demander_from['address']}")
                run3.font.size = Pt(14)
                run3.font.name = 'Times New Roman'

        p = self.doc.add_paragraph()
        run = p.add_run("")
        run.font.size = Pt(6)

        self.add_title("Досудебная претензия")

        if data.get('demand_basis'):
            self.add_section_header("1. Основание претензии")
            self.add_body_text(data['demand_basis'])

        if data.get('description'):
            self.add_section_header("2. Описание нарушения")
            self.add_body_text(data['description'])

        self.add_section_header("3. Требования")
        if data.get('claims'):
            claims = data['claims'] if isinstance(data['claims'], list) else data['claims'].split('\n')
            for i, claim in enumerate(claims, 1):
                if claim.strip():
                    self.add_list_item(claim.strip(), number=str(i))
        else:
            self.add_list_item("Устранить допущенные нарушения в срок.")
            self.add_list_item("Компенсировать причинённый ущерб.")

        deadline = data.get('demand_deadline', '10 календарных дней')
        self.add_section_header("4. Срок исполнения")
        self.add_body_text(f"Требую исполнить указанные требования в срок: {deadline}")
        self.add_body_text("В случае неисполнения требований буду вынужден обратиться в суд с отнесением всех судебных расходов на ваш счёт.")

        if demander_from.get('name'):
            self.add_signature_block(demander_from['name'])

        # Добавляем водяной знак перед сохранением
        self._add_watermark()

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_from_plain_text(self, title: str, content: str) -> bytes:
        """
        Сгенерировать .docx из обычного текста (для исправленного договора).
        Разбивает текст на параграфы по пустым строкам.
        """
        self.doc = Document()
        self._setup_styles()

        # Заголовок
        if title:
            self.add_title(title)

        # Разбиваем текст на параграфы
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Заголовки (короткие строки без точки в конце)
            if len(para_text) < 80 and not para_text.endswith('.') and para_text.isupper():
                self.add_title(para_text)
            elif len(para_text) < 100 and not para_text.endswith('.') and para_text[0].isupper():
                self.add_subtitle(para_text)
            else:
                # Обычный текст - разбиваем по строкам
                lines = para_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    # Списки
                    if line.startswith(('- ', '• ', '● ', '✓ ', '✔ ')):
                        self.add_list_item(line[2:].strip())
                    elif line[0].isdigit() and '. ' in line[:5]:
                        self.add_list_item(line.split('. ', 1)[1].strip(), number=line[0])
                    else:
                        self.add_body_text(line)

            # Пустая строка между параграфами
            p = self.doc.add_paragraph()
            run = p.add_run("")
            run.font.size = Pt(6)

        # Добавляем водяной знак перед сохранением
        self._add_watermark()

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# Singleton
docx_generator = DocxGenerator()

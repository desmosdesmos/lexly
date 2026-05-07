from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel, Field
import logging
import os
import uuid
import json
import io
from pathlib import Path

from app.database import get_db
from app.models.user import User
from app.models.contract_review import ContractReview
from app.middleware.auth import get_current_user
from app.services.ai_service import ai_service
from app.services.limit_service import limit_service
from app.services.docx_generator import docx_generator

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/contracts", tags=["Договоры"])

# Директория для загрузки файлов
UPLOAD_DIR = Path("./uploads/contracts")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 МБ


async def extract_text_from_file(file_path: Path) -> str:
    """Извлечение текста из файла."""
    ext = file_path.suffix.lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Ошибка извлечения текста из PDF: {str(e)}",
            )

    elif ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Ошибка извлечения текста из документа: {str(e)}",
            )

    elif ext == ".doc":
        # .doc файлы (Word 97-2003) требуют специальной обработки
        try:
            # Пробуем через subprocess libreoffice (если установлен)
            import subprocess
            import tempfile
            
            # Конвертируем .doc в .docx через libreoffice
            output_dir = tempfile.gettempdir()
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', output_dir, str(file_path)],
                capture_output=True, text=True, timeout=30
            )
            
            if result.returncode == 0:
                docx_path = str(file_path).replace('.doc', '.docx')
                if not os.path.exists(docx_path):
                    docx_path = os.path.join(output_dir, file_path.stem + '.docx')
                
                if os.path.exists(docx_path):
                    from docx import Document as DocxDocument
                    doc = DocxDocument(docx_path)
                    text = "\n".join([p.text for p in doc.paragraphs])
                    os.remove(docx_path)  # Удаляем временный файл
                    return text
        except subprocess.TimeoutExpired:
            pass
        except Exception:
            pass
        
        # Если libreoffice не установлен - пробуем через olefile (сырой текст)
        try:
            import olefile
            ole = olefile.OleFileIO(str(file_path))
            text_parts = []
            for stream in ole.listdir():
                if stream and len(stream) > 0:
                    try:
                        data = ole.openstream(stream).read()
                        text = data.decode('utf-8', errors='ignore')
                        if text.strip():
                            text_parts.append(text)
                    except:
                        pass
            ole.close()
            if text_parts:
                return '\n'.join(text_parts[:10])  # Ограничиваем
        except Exception:
            pass
        
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Не удалось извлечь текст из .doc файла. Конвертируйте файл в .docx или PDF.",
        )

    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Неподдерживаемый формат файла: {ext}",
        )


@router.post(
    "/review",
    summary="Проверка договора через AI",
    status_code=status.HTTP_200_OK,
)
async def review_contract(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Загрузить договор для AI-анализа на риски.
    
    Поддерживаемые форматы: PDF, DOC, DOCX, TXT
    Максимальный размер: 10 МБ
    """
    # Проверка размера файла
    file_content = await file.read()
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Размер файла не должен превышать 10 МБ",
        )
    
    # Проверка формата
    file_ext = Path(file.filename).suffix.lower() if file.filename else ""
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Неподдерживаемый формат. Допускаются: {', '.join(ALLOWED_EXTENSIONS)}",
        )
    
    # Проверка лимитов
    can_use, limit_info = await limit_service.check_contract_limit(
        user_id=current_user.id,
        db=db,
    )

    if not can_use:
        raise HTTPException(
            status_code=status.HTTP_402_PAYMENT_REQUIRED,
            detail={
                "type": "limit_exceeded",
                "resource": "contracts",
                "message": f"Вы достигли лимита проверки договоров ({limit_info['used']}/{limit_info['max']}). Перейдите на Pro.",
                "limit_info": limit_info,
            },
        )
    
    # Сохранение файла
    file_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{file_id}{file_ext}"
    
    with open(file_path, "wb") as f:
        f.write(file_content)
    
    try:
        # Извлечение текста
        extracted_text = await extract_text_from_file(file_path)
        
        if not extracted_text or len(extracted_text.strip()) < 100:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Не удалось извлечь достаточно текста из документа",
            )
        
        # AI анализ
        analysis_result = await ai_service.review_contract(extracted_text)
        
        # Определение уровня риска
        risks = analysis_result.get("risks", [])
        if not risks:
            risk_level = "low"
        else:
            severities = [r.get("severity", "low") for r in risks]
            if "critical" in severities:
                risk_level = "critical"
            elif "high" in severities:
                risk_level = "high"
            elif "medium" in severities:
                risk_level = "medium"
            else:
                risk_level = "low"
        
        # Создание записи в БД
        review = ContractReview(
            user_id=current_user.id,
            original_file_name=file.filename or "unknown",
            file_path=str(file_path),
            extracted_text=extracted_text[:10000],  # Ограничиваем длину
            analysis_result=json.dumps(analysis_result, ensure_ascii=False),  # Конвертируем dict в JSON строку
            risks=json.dumps(risks, ensure_ascii=False),  # Конвертируем list в JSON строку
            recommendations=json.dumps(analysis_result.get("recommendations", []), ensure_ascii=False),  # Конвертируем list в JSON строку
            risk_level=risk_level,
        )
        
        db.add(review)
        await db.commit()
        await db.refresh(review)

        # Инкремент использования
        await limit_service.increment_contracts(current_user.id, db)
        
        logger.info(f"Contract reviewed: {review.id} for user {current_user.id}")
        
        return {
            "id": str(review.id),
            "original_file_name": review.original_file_name,
            "status": "completed",
            "risk_level": review.risk_level,
            "analysis": json.loads(review.analysis_result) if isinstance(review.analysis_result, str) else review.analysis_result,
            "created_at": str(review.created_at),
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Contract review error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка анализа договора: {str(e)}",
        )


@router.get(
    "",
    summary="Список проверок договоров",
)
async def list_reviews(
    page: int = 1,
    limit: int = 20,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Получить список всех проверок пользователя."""
    offset = (page - 1) * limit
    
    # Общее количество
    count_result = await db.execute(
        select(func.count()).select_from(ContractReview).where(ContractReview.user_id == current_user.id)
    )
    total = count_result.scalar() or 0
    
    # Список проверок
    reviews_result = await db.execute(
        select(ContractReview)
        .where(ContractReview.user_id == current_user.id)
        .order_by(ContractReview.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    reviews = reviews_result.scalars().all()
    
    items = [
        {
            "id": str(review.id),
            "original_file_name": review.original_file_name,
            "status": "completed",
            "risk_level": review.risk_level,
            "created_at": str(review.created_at),
        }
        for review in reviews
    ]
    
    return {
        "items": items,
        "total": total,
        "page": page,
        "limit": limit,
    }


@router.get(
    "/{review_id}",
    summary="Получить результат проверки",
)
async def get_review(
    review_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Получить конкретный результат проверки."""
    result = await db.execute(
        select(ContractReview).where(
            ContractReview.id == review_id,
            ContractReview.user_id == current_user.id,
        )
    )
    review = result.scalar_one_or_none()
    
    if not review:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Проверка не найдена",
        )
    
    return {
        "id": str(review.id),
        "original_file_name": review.original_file_name,
        "status": "completed",
        "risk_level": review.risk_level,
        "analysis": json.loads(review.analysis_result) if isinstance(review.analysis_result, str) else review.analysis_result,
        "created_at": str(review.created_at),
        "completed_at": str(review.completed_at) if review.completed_at else None,
    }


@router.delete(
    "/{review_id}",
    summary="Удалить проверку",
    status_code=status.HTTP_200_OK,
)
async def delete_review(
    review_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Удалить результат проверки."""
    result = await db.execute(
        select(ContractReview).where(
            ContractReview.id == review_id,
            ContractReview.user_id == current_user.id,
        )
    )
    review = result.scalar_one_or_none()
    
    if not review:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Проверка не найдена",
        )
    
    # Удаление файла
    if review.file_path and os.path.exists(review.file_path):
        os.remove(review.file_path)
    
    await db.delete(review)
    await db.commit()
    
    return {"message": "Проверка успешно удалена"}


# ---- AI-корректировка договора ----

class ContractFixRequest(BaseModel):
    review_id: str = Field(..., description="ID проверки договора")
    risks_to_fix: Optional[List[str]] = Field(None, description="Список ID рисков для исправления (если None - все)")


@router.post(
    "/{review_id}/fix",
    summary="AI-корректировка договора (исправление найденных рисков)",
)
async def fix_contract(
    review_id: str,
    request: ContractFixRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    AI анализирует найденные риски и предлагает исправленную версию договора.
    """
    # Находим проверку
    result = await db.execute(
        select(ContractReview).where(
            ContractReview.id == review_id,
            ContractReview.user_id == current_user.id,
        )
    )
    review = result.scalar_one_or_none()

    if not review:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Проверка не найдена",
        )

    # Получаем текст договора и риски
    contract_text = review.extracted_text
    risks = json.loads(review.risks) if isinstance(review.risks, str) else review.risks
    analysis = json.loads(review.analysis_result) if isinstance(review.analysis_result, str) else review.analysis_result

    # Фильтруем риски если указаны конкретные
    risks_to_fix = request.risks_to_fix or []
    filtered_risks = [r for r in risks if not risks_to_fix or r.get('id') in risks_to_fix]

    # Формируем промпт для AI
    risks_description = "\n".join([
        f"- {r.get('clause', 'Пункт')}: {r.get('text', '')} → {r.get('recommendation', '')}"
        for r in filtered_risks
    ])

    system_prompt = """Ты — AI-юрист, специализирующийся на исправлении договоров.
Твоя задача — исправить конкретные проблемы в тексте договора, сохраняя остальной текст без изменений.

КРИТИЧЕСКИЕ ПРАВИЛА:
1. Исправляй ТОЛЬКО указанные проблемы
2. Сохраняй весь остальной текст договора БЕЗ ИЗМЕНЕНИЙ
3. НЕ добавляй новые разделы, которых не было
4. НЕ удаляй существенные разделы без указания
5. Используй правильный юридический язык
6. Верни ПОЛНЫЙ текст договора с исправлениями, а не только изменённые части
7. В начале ответа добавь краткий список изменений (3-5 пунктов)"""

    user_prompt = f"""Исправь следующие проблемы в договоре:

{risks_description}

ПОЛНЫЙ ТЕКСТ ДОГОВОРА:
---
{contract_text}
---

Верни полный исправленный договор."""

    try:
        fixed_text = await ai_service.generate(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            temperature=0.2,
            max_tokens=8192,
        )

        # Обновляем запись в БД
        review.fixed_content = fixed_text
        review.status = "fixed"
        await db.commit()

        return {
            "review_id": str(review.id),
            "status": "fixed",
            "fixed_content": fixed_text,
            "fixed_risks_count": len(filtered_risks),
        }

    except Exception as e:
        logger.error(f"Ошибка AI-корректировки договора: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка корректировки: {str(e)}",
        )


@router.get(
    "/{review_id}/download-fixed",
    summary="Скачать исправленный договор в .docx",
)
async def download_fixed_contract(
    review_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Скачать исправленный AI договор в формате .docx."""
    result = await db.execute(
        select(ContractReview).where(
            ContractReview.id == review_id,
            ContractReview.user_id == current_user.id,
        )
    )
    review = result.scalar_one_or_none()

    if not review:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Проверка не найдена",
        )

    if not review.fixed_content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Договор ещё не был исправлен AI. Сначала запустите корректировку.",
        )

    # Генерируем .docx
    try:
        docx_bytes = docx_generator.generate_from_plain_text(
            title=f"ИСПРАВЛЕННЫЙ ДОГОВОР\n{review.original_file_name}",
            content=review.fixed_content,
        )
    except Exception as e:
        logger.error(f"Ошибка генерации .docx: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка генерации файла: {str(e)}",
        )

    # Имя файла
    safe_name = f"fixed_{review.original_file_name.rsplit('.', 1)[0] if '.' in review.original_file_name else review.original_file_name}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
    )
